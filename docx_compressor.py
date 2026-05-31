import io
import os
import shutil
import zipfile


SUPPORTED_DOCX_QUALITIES = {
    "low": {"jpeg_quality": 40, "max_dim": 800},
    "medium": {"jpeg_quality": 65, "max_dim": 1200},
    "high": {"jpeg_quality": 80, "max_dim": 1600},
}


class CorruptDocxError(Exception):
    pass


def _optimize_image(data: bytes, extension: str, jpeg_quality: int, max_dim: int) -> bytes | None:
    try:
        from PIL import Image
    except ModuleNotFoundError as exc:
        raise RuntimeError(
            "Missing DOCX image dependency. Install Pillow from requirements.txt."
        ) from exc

    try:
        image = Image.open(io.BytesIO(data))
        width, height = image.size
        if max(width, height) > max_dim:
            scale = max_dim / max(width, height)
            image = image.resize((int(width * scale), int(height * scale)), Image.LANCZOS)

        output = io.BytesIO()
        if extension in (".jpg", ".jpeg"):
            image.convert("RGB").save(
                output,
                format="JPEG",
                quality=jpeg_quality,
                optimize=True,
            )
        elif extension == ".png":
            image.save(output, format="PNG", optimize=True)
        else:
            return None
        return output.getvalue()
    except Exception:
        return None


def compress_docx(input_path: str, output_path: str, quality: str = "medium") -> dict:
    try:
        from docx import Document
        from docx.opc.exceptions import PackageNotFoundError
    except ModuleNotFoundError as exc:
        raise RuntimeError(
            "Missing DOCX dependency. Install python-docx from requirements.txt."
        ) from exc

    settings = SUPPORTED_DOCX_QUALITIES.get(quality, SUPPORTED_DOCX_QUALITIES["medium"])
    jpeg_quality = settings["jpeg_quality"]
    max_dim = settings["max_dim"]

    try:
        document = Document(input_path)
        paragraph_count = len(document.paragraphs)
        table_count = len(document.tables)
    except (PackageNotFoundError, zipfile.BadZipFile, KeyError, ValueError) as exc:
        raise CorruptDocxError() from exc

    tmp_path = output_path + ".tmp.zip"
    image_count = 0

    try:
        with zipfile.ZipFile(input_path, "r") as zin, zipfile.ZipFile(
            tmp_path,
            "w",
            compression=zipfile.ZIP_DEFLATED,
            compresslevel=9,
        ) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                extension = os.path.splitext(item.filename)[1].lower()

                if item.filename.startswith("word/media/") and extension in (".png", ".jpg", ".jpeg"):
                    optimized = _optimize_image(data, extension, jpeg_quality, max_dim)
                    if optimized:
                        data = optimized
                        image_count += 1

                zout.writestr(item, data)
    except (zipfile.BadZipFile, KeyError, ValueError) as exc:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)
        raise CorruptDocxError() from exc

    shutil.move(tmp_path, output_path)

    original_size = os.path.getsize(input_path)
    compressed_size = os.path.getsize(output_path)
    space_saved = round((1 - compressed_size / original_size) * 100, 2) if original_size else 0
    ratio = round(original_size / compressed_size, 4) if compressed_size else 1.0

    return {
        "paragraph_count": paragraph_count,
        "image_count": image_count,
        "table_count": table_count,
        "original_size": original_size,
        "compressed_size": compressed_size,
        "space_saved_percent": space_saved,
        "compression_ratio": ratio,
        "total_chars": 0,
        "total_bits": 0,
    }
